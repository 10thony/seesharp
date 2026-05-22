#!/usr/bin/env python3
"""Generate a Survey123 feature report template focused on notes, issues, and photos."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from generate_survey123_report_template import (
        SurveyField,
        SurveyForm,
        SurveyGroup,
        SurveyRepeat,
        base_type,
        clean_text,
        field_placeholder,
        parse_xlsform,
    )
except ModuleNotFoundError:
    from tools.generate_survey123_report_template import (
        SurveyField,
        SurveyForm,
        SurveyGroup,
        SurveyRepeat,
        base_type,
        clean_text,
        field_placeholder,
        parse_xlsform,
    )

BODY_FONT = "Aptos"

REPORT_BLUE_HEX = "0A1F62"
REPORT_BLUE = RGBColor.from_string(REPORT_BLUE_HEX)
LIGHT_BLUE_FILL = "EAF3F6"
SECTION_FILL = REPORT_BLUE_HEX
CARD_COLUMNS = 2

TEXT_MATCH_TERMS = (
    "note",
    "notes",
    "issue",
    "issues",
    "finding",
    "findings",
)
IMAGE_MATCH_TERMS = ("photo", "photos", "image", "images")


def set_shading(cell_or_paragraph, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    if hasattr(cell_or_paragraph, "_tc"):
        cell_or_paragraph._tc.get_or_add_tcPr().append(shading)
    else:
        cell_or_paragraph._p.get_or_add_pPr().append(shading)


def set_cell_width(cell, width: Inches) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(
    cell, *, top: int = 120, start: int = 140, bottom: int = 120, end: int = 140
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)

    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D8DEE3", size: str = "5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_text(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_section_band(doc: Document, text: str, *, major: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_shading(paragraph, SECTION_FILL if major else LIGHT_BLUE_FILL)
    add_text(
        paragraph,
        clean_text(text),
        bold=True,
        size=12 if major else 10,
        color=RGBColor(0xFF, 0xFF, 0xFF) if major else REPORT_BLUE,
    )


def is_report_field(field: SurveyField) -> bool:
    field_type = base_type(field.type)
    haystack = f"{field.name} {field.label} {field.type}".lower()
    if field_type == "note":
        return False
    if field_type == "image":
        return True
    if field_type == "file":
        return any(term in haystack for term in IMAGE_MATCH_TERMS)
    return any(term in haystack for term in TEXT_MATCH_TERMS + IMAGE_MATCH_TERMS)


def report_fields(fields: list[SurveyField]) -> list[SurveyField]:
    return [field for field in fields if is_report_field(field)]


def strip_outer_parens(expression: str) -> str:
    expression = expression.strip()
    while expression.startswith("(") and expression.endswith(")"):
        depth = 0
        wraps_entire_expression = True
        for index, char in enumerate(expression):
            if char == "(":
                depth += 1
            elif char == ")":
                depth -= 1
                if depth == 0 and index != len(expression) - 1:
                    wraps_entire_expression = False
                    break
        if not wraps_entire_expression:
            break
        expression = expression[1:-1].strip()
    return expression


def report_condition_from_xlsform_expression(expression: str) -> str:
    expression = strip_outer_parens(expression)
    if not expression or expression.lower() in {"yes", "true"}:
        return ""

    if_match = re.fullmatch(
        r"if\((?P<condition>.+),\s*['\"]yes['\"]\s*,\s*['\"]{0,1}['\"]\s*\)",
        expression,
        flags=re.I,
    )
    if if_match:
        expression = if_match.group("condition").strip()

    expression = strip_outer_parens(expression)
    expression = re.sub(r"\$\{([A-Za-z_][A-Za-z0-9_]*)\}", r"\1", expression)
    expression = expression.replace("'", '"')
    expression = re.sub(r"(?<![<>=!])=(?!=)", "==", expression)
    return expression


def field_visibility_condition(field: SurveyField) -> str:
    return report_condition_from_xlsform_expression(field.visible or field.relevant)


def conditional_open(field: SurveyField) -> str:
    condition = field_visibility_condition(field)
    return f"${{if {condition}}}" if condition else ""


def conditional_close(field: SurveyField) -> str:
    return "${/}" if field_visibility_condition(field) else ""


def group_has_report_content(group: SurveyGroup) -> bool:
    if report_fields(group.fields):
        return True
    if any(group_has_report_content(child) for child in group.groups):
        return True
    return any(repeat_has_report_content(repeat) for repeat in group.repeats)


def repeat_has_report_content(repeat: SurveyRepeat) -> bool:
    if report_fields(repeat.fields):
        return True
    return any(group_has_report_content(group) for group in repeat.groups)


def add_field_card(cell, field: SurveyField) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=150, bottom=150)

    label = cell.paragraphs[0]
    add_text(
        label,
        f"{conditional_open(field)}{field.label}",
        bold=True,
        size=8,
        color=REPORT_BLUE,
    )
    label.paragraph_format.space_after = Pt(3)

    value = cell.add_paragraph()
    add_text(value, f"{field_placeholder(field)}{conditional_close(field)}", size=9)
    value.paragraph_format.space_after = Pt(0)


def add_empty_card(cell) -> None:
    cell.text = ""
    set_shading(cell, "F7F9FA")
    set_cell_margins(cell)


def add_field_cards(doc: Document, fields: list[SurveyField]) -> None:
    fields = report_fields(fields)
    if not fields:
        return

    row_count = (len(fields) + CARD_COLUMNS - 1) // CARD_COLUMNS
    table = doc.add_table(rows=row_count, cols=CARD_COLUMNS)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="E1E6EA", size="4")

    for row_index in range(row_count):
        for column_index in range(CARD_COLUMNS):
            field_index = row_index * CARD_COLUMNS + column_index
            cell = table.rows[row_index].cells[column_index]
            set_cell_width(cell, Inches(3.45))
            if field_index < len(fields):
                add_field_card(cell, fields[field_index])
            else:
                add_empty_card(cell)
    doc.add_paragraph()


def repeat_rows(repeat: SurveyRepeat) -> list[tuple[str, SurveyField | None]]:
    rows: list[tuple[str, SurveyField | None]] = []
    fields = report_fields(repeat.fields)
    if fields:
        rows.append(("Details", None))
        rows.extend((field.label, field) for field in fields)

    def collect_group(group: SurveyGroup, prefix: str = "") -> None:
        label = group.label or group.name
        section_label = f"{prefix} - {label}" if prefix and label else label or prefix
        group_fields = report_fields(group.fields)
        if group_fields:
            rows.append((section_label or "Details", None))
            rows.extend((field.label, field) for field in group_fields)
        for child in group.groups:
            collect_group(child, section_label)

    for group in repeat.groups:
        collect_group(group)
    return rows


def add_repeat_table(doc: Document, repeat: SurveyRepeat) -> None:
    rows = repeat_rows(repeat)
    if not rows:
        return

    start = doc.add_paragraph()
    add_text(start, f"${{# {repeat.name}}}", bold=True, size=8, color=REPORT_BLUE)
    add_section_band(doc, repeat.label or repeat.name, major=False)

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    field_header, response_header = table.rows[0].cells
    for cell in (field_header, response_header):
        cell.text = ""
        set_shading(cell, LIGHT_BLUE_FILL)
        set_cell_margins(cell)
    set_cell_width(field_header, Inches(2.45))
    set_cell_width(response_header, Inches(4.55))
    add_text(field_header.paragraphs[0], "Field", bold=True, size=8, color=REPORT_BLUE)
    add_text(response_header.paragraphs[0], "Response", bold=True, size=8, color=REPORT_BLUE)

    for row_index, (label, field) in enumerate(rows, start=1):
        label_cell, value_cell = table.rows[row_index].cells
        label_cell.text = ""
        value_cell.text = ""
        set_cell_width(label_cell, Inches(2.45))
        set_cell_width(value_cell, Inches(4.55))
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)

        if field is None:
            merged = label_cell.merge(value_cell)
            merged.text = ""
            set_shading(merged, "F0F4F6")
            set_cell_margins(merged, top=90, bottom=90)
            add_text(merged.paragraphs[0], label, bold=True, size=8, color=REPORT_BLUE)
            continue

        add_text(
            label_cell.paragraphs[0],
            f"{conditional_open(field)}{label}",
            bold=True,
            size=8,
        )
        add_text(
            value_cell.paragraphs[0],
            f"{field_placeholder(field)}{conditional_close(field)}",
            size=9,
        )

    end = doc.add_paragraph()
    add_text(end, "${/}", bold=True, size=8, color=REPORT_BLUE)


def add_cover(doc: Document, form: SurveyForm) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(title, form.title, bold=True, size=22, color=REPORT_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(subtitle, "Notes, Issues, and Photos Report Template", bold=True, size=10)

    table = doc.add_table(rows=2 if form.version else 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    rows = [("Report Date", '${$date | format:"MM/DD/YYYY HH:mm"}')]
    if form.version:
        rows.append(("XLSForm Version", form.version))
    for index, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[index].cells
        label_cell.text = ""
        value_cell.text = ""
        set_cell_width(label_cell, Inches(2.0))
        set_cell_width(value_cell, Inches(5.0))
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        set_shading(label_cell, LIGHT_BLUE_FILL)
        add_text(label_cell.paragraphs[0], label, bold=True, size=8, color=REPORT_BLUE)
        add_text(value_cell.paragraphs[0], value, size=8)
    doc.add_paragraph()


def render_group(doc: Document, group: SurveyGroup, depth: int = 1) -> None:
    if not group_has_report_content(group):
        return

    add_section_band(doc, group.label or group.name, major=depth == 1)
    add_field_cards(doc, group.fields)

    for repeat in group.repeats:
        add_repeat_table(doc, repeat)
    for child in group.groups:
        render_group(doc, child, depth + 1)


def render_form(doc: Document, form: SurveyForm) -> None:
    for item in form.items:
        if isinstance(item, SurveyField):
            add_field_cards(doc, [item])
        elif isinstance(item, SurveyRepeat):
            add_repeat_table(doc, item)
        else:
            render_group(doc, item)


def generate_report_template(xlsx_path: Path, output_path: Path) -> Path:
    form = parse_xlsform(xlsx_path)
    document = Document()
    add_cover(document, form)
    render_form(document, form)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a Survey123 notes/issues/photos feature report template."
    )
    parser.add_argument("--xlsx", required=True, type=Path, help="Path to the XLSForm.")
    parser.add_argument(
        "--output",
        required=True,
        type=Path,
        help="Output .docx path for the generated report template.",
    )
    args = parser.parse_args()

    if not args.xlsx.exists():
        raise SystemExit(f"XLSForm not found: {args.xlsx}")

    output = generate_report_template(args.xlsx, args.output)
    print(f"Wrote report template: {output}")


if __name__ == "__main__":
    main()
