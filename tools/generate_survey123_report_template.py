#!/usr/bin/env python3
"""Generate an ArcGIS Survey123 feature report template from an XLSForm."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from html import unescape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook

BODY_FONT = "Aptos"
TITLE_COLOR = RGBColor(0x1F, 0x4E, 0x5F)
SECTION_FILL = "1F4E5F"
SUBSECTION_FILL = "D9EAF0"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
CARD_COLUMNS = 2

SKIP_TYPES = {
    "note",
    "calculate",
    "begin group",
    "end group",
    "begin repeat",
    "end repeat",
}

MEDIA_TYPES = {"image", "audio", "file"}
MAP_TYPES = {"geopoint", "geotrace", "geoshape"}
DATE_TYPES = {"date", "datetime"}


@dataclass
class Choice:
    name: str
    label: str


@dataclass
class SurveyField:
    type: str
    name: str
    label: str
    appearance: str = ""
    relevant: str = ""
    visible: str = ""
    list_name: str | None = None
    choices: list[Choice] = field(default_factory=list)


@dataclass
class SurveyGroup:
    name: str
    label: str
    fields: list[SurveyField] = field(default_factory=list)
    groups: list["SurveyGroup"] = field(default_factory=list)
    repeats: list["SurveyRepeat"] = field(default_factory=list)


@dataclass
class SurveyRepeat:
    name: str
    label: str
    fields: list[SurveyField] = field(default_factory=list)
    groups: list[SurveyGroup] = field(default_factory=list)


@dataclass
class SurveyForm:
    title: str
    version: str
    items: list[SurveyField | SurveyGroup | SurveyRepeat] = field(default_factory=list)


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = unescape(str(value))
    text = re.sub(r"<br\s*/?>", " ", text, flags=re.I)
    text = re.sub(r"</(?:p|h\d|div|li)>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("\xa0", " ")
    return re.sub(r"\s+", " ", text).strip()


def base_type(type_name: str) -> str:
    return type_name.strip().split(maxsplit=1)[0].lower()


def choice_list_from_type(type_name: str) -> str | None:
    match = re.match(
        r"select_(?:one|multiple)(?:_from_file)?\s+(.+)",
        type_name.strip(),
        flags=re.I,
    )
    return match.group(1).strip() if match else None


def load_choices(workbook) -> dict[str, list[Choice]]:
    if "choices" not in workbook.sheetnames:
        return {}

    worksheet = workbook["choices"]
    rows = worksheet.iter_rows(values_only=True)
    header = next(rows, None)
    if not header:
        return {}

    columns = {str(name).strip(): i for i, name in enumerate(header) if name}
    list_idx = columns.get("list_name", 0)
    name_idx = columns.get("name", 1)
    label_idx = columns.get("label", 2)

    choices: dict[str, list[Choice]] = {}
    for row in rows:
        if not row or list_idx >= len(row) or not row[list_idx]:
            continue
        list_name = str(row[list_idx]).strip()
        choice_name = row[name_idx] if name_idx < len(row) else None
        if choice_name is None or str(choice_name).strip() == "":
            continue
        choice_label = row[label_idx] if label_idx < len(row) else choice_name
        choices.setdefault(list_name, []).append(
            Choice(str(choice_name).strip(), clean_text(choice_label) or str(choice_name))
        )
    return choices


def load_settings(workbook) -> tuple[str, str]:
    if "settings" not in workbook.sheetnames:
        return "Survey Feature Report", ""

    worksheet = workbook["settings"]
    rows = list(worksheet.iter_rows(values_only=True))
    if len(rows) < 2:
        return "Survey Feature Report", ""

    header = [str(value).strip() if value else "" for value in rows[0]]
    values = dict(zip(header, rows[1]))
    title = clean_text(values.get("form_title")) or "Survey Feature Report"
    version = str(values.get("version") or "").strip()
    return title, version


def row_value(row: tuple[Any, ...], columns: dict[str, int], column_name: str) -> Any:
    index = columns.get(column_name)
    if index is None or index >= len(row):
        return None
    return row[index]


def should_skip_field(row: dict[str, Any]) -> bool:
    field_type = str(row.get("type") or "").strip().lower()
    field_name = str(row.get("name") or "").strip()
    if not field_name or field_name == "list_name":
        return True
    if field_type in SKIP_TYPES or base_type(field_type) in SKIP_TYPES:
        return True
    if str(row.get("visible") or "").strip().lower() == "no":
        return True
    if str(row.get("field_type") or "").strip().lower() == "null":
        return True
    return False


def make_field(row: dict[str, Any], choices: dict[str, list[Choice]]) -> SurveyField:
    type_name = str(row["type"]).strip()
    list_name = choice_list_from_type(type_name)
    label = clean_text(row.get("label")) or clean_text(row.get("alias"))
    return SurveyField(
        type=type_name,
        name=str(row["name"]).strip(),
        label=label or str(row["name"]).strip(),
        appearance=str(row.get("appearance") or "").strip(),
        relevant=str(row.get("relevant") or "").strip(),
        visible=str(row.get("visible") or "").strip(),
        list_name=list_name,
        choices=choices.get(list_name or "", []),
    )


def parse_xlsform(xlsx_path: Path) -> SurveyForm:
    workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
    try:
        if "survey" not in workbook.sheetnames:
            raise ValueError(f"Workbook does not contain a survey sheet: {xlsx_path}")

        title, version = load_settings(workbook)
        choices = load_choices(workbook)
        worksheet = workbook["survey"]

        header_row = next(worksheet.iter_rows(max_row=1, values_only=True))
        header = [str(value).strip() if value else "" for value in header_row]
        columns = {name: index for index, name in enumerate(header) if name}

        root_items: list[SurveyField | SurveyGroup | SurveyRepeat] = []
        stack: list[SurveyGroup | SurveyRepeat] = []

        def current_parent() -> SurveyGroup | SurveyRepeat | None:
            return stack[-1] if stack else None

        def append_item(item: SurveyField | SurveyGroup | SurveyRepeat) -> None:
            parent = current_parent()
            if parent is None:
                root_items.append(item)
            elif isinstance(parent, SurveyGroup):
                if isinstance(item, SurveyField):
                    parent.fields.append(item)
                elif isinstance(item, SurveyRepeat):
                    parent.repeats.append(item)
                else:
                    parent.groups.append(item)
            elif isinstance(item, SurveyField):
                parent.fields.append(item)
            elif isinstance(item, SurveyGroup):
                parent.groups.append(item)
            else:
                root_items.append(item)

        for row in worksheet.iter_rows(min_row=2, values_only=True):
            mapped = {
                "type": row_value(row, columns, "type"),
                "name": row_value(row, columns, "name"),
                "label": row_value(row, columns, "label"),
                "appearance": row_value(row, columns, "appearance"),
                "relevant": row_value(row, columns, "relevant"),
                "visible": row_value(row, columns, "body::esri:visible"),
                "field_type": row_value(row, columns, "bind::esri:fieldType"),
                "alias": row_value(row, columns, "bind::esri:fieldAlias"),
            }
            row_type = str(mapped["type"] or "").strip()
            normalized_type = row_type.lower()
            if not row_type:
                continue

            if normalized_type == "begin group":
                group = SurveyGroup(
                    name=str(mapped.get("name") or "").strip(),
                    label=clean_text(mapped.get("label"))
                    or str(mapped.get("name") or "").strip(),
                )
                append_item(group)
                stack.append(group)
                continue

            if normalized_type == "end group":
                while stack and not isinstance(stack[-1], SurveyGroup):
                    stack.pop()
                if stack:
                    stack.pop()
                continue

            if normalized_type == "begin repeat":
                repeat = SurveyRepeat(
                    name=str(mapped.get("name") or "").strip(),
                    label=clean_text(mapped.get("label"))
                    or str(mapped.get("name") or "").strip(),
                )
                append_item(repeat)
                stack.append(repeat)
                continue

            if normalized_type == "end repeat":
                while stack and not isinstance(stack[-1], SurveyRepeat):
                    stack.pop()
                if stack:
                    stack.pop()
                continue

            if not should_skip_field(mapped):
                append_item(make_field(mapped, choices))

        return SurveyForm(title=title, version=version, items=root_items)
    finally:
        workbook.close()


def choice_list_name(field_item: SurveyField) -> str | None:
    return field_item.list_name or choice_list_from_type(field_item.type)


def is_yes_only(field_item: SurveyField) -> bool:
    list_name = choice_list_name(field_item)
    if list_name == "yes_only":
        return True
    choice_names = {choice.name.lower() for choice in field_item.choices}
    return choice_names == {"yes"}


def is_yes_no_family(field_item: SurveyField) -> bool:
    list_name = choice_list_name(field_item)
    return list_name in {"yes_no", "yes_no_unk", "yes_no_na"}


def checkbox(name: str, value: str, label: str | None = None) -> str:
    text = f'${{{name} | checked:"{value}"}}'
    return f"{text} {label}" if label else text


def field_placeholder(field_item: SurveyField, include_choice_labels: bool = True) -> str:
    type_name = base_type(field_item.type)
    name = field_item.name

    if is_yes_only(field_item):
        return checkbox(name, "yes")

    if is_yes_no_family(field_item):
        parts = [checkbox(name, "yes", "Yes"), checkbox(name, "no", "No")]
        list_name = choice_list_name(field_item)
        if list_name == "yes_no_unk":
            parts.append(checkbox(name, "unknown", "Unknown"))
        elif list_name == "yes_no_na":
            parts.append(checkbox(name, "n/a", "N/A"))
        return "  ".join(parts)

    if type_name.startswith("select_one") and field_item.choices and len(field_item.choices) <= 10:
        return "  ".join(
            checkbox(name, choice.name, choice.label if include_choice_labels else None)
            for choice in field_item.choices
        )

    if type_name.startswith("select_multiple"):
        if field_item.choices and len(field_item.choices) <= 10:
            return "  ".join(
                checkbox(name, choice.name, choice.label if include_choice_labels else None)
                for choice in field_item.choices
            )
        return f'${{{name} | appearance:"bullets"}}'

    if type_name in DATE_TYPES:
        return f'${{{name} | format:"MM/DD/YYYY"}}'

    if type_name in MEDIA_TYPES:
        return (
            f"${{if {name}}}"
            f"${{# {name}}}"
            f'${{$file | getValue:"" | size:0:0:600:450}}'
            f"${{/}}"
            f"${{/}}"
        )

    if type_name in MAP_TYPES:
        return f"${{{name}}}"

    if type_name == "text" and "multiline" in field_item.appearance.lower():
        return f'${{{name} | appearance:"multiline"}}'

    return f"${{{name}}}"


def set_shading(cell_or_paragraph: Any, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    if hasattr(cell_or_paragraph, "_tc"):
        cell_or_paragraph._tc.get_or_add_tcPr().append(shading)
    else:
        cell_or_paragraph._p.get_or_add_pPr().append(shading)


def set_cell_width(cell: Any, width: Inches) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(
    cell: Any, *, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)

    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table: Any, color: str = "D8DEE3", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_text(
    paragraph: Any,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    color: RGBColor | None = None,
) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_section_band(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_shading(paragraph, SECTION_FILL if level == 1 else SUBSECTION_FILL)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(12 if level == 1 else 10)
    run.font.color.rgb = HEADER_TEXT if level == 1 else TITLE_COLOR


def add_label_value_table(
    doc: Document, rows: list[tuple[str, str]], label_width: float = 2.6
) -> None:
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for index, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[index].cells
        label_cell.text = ""
        value_cell.text = ""
        set_cell_width(label_cell, Inches(label_width))
        set_cell_width(value_cell, Inches(7.0 - label_width))
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_shading(label_cell, "EAF3F6")
        add_text(label_cell.paragraphs[0], label, bold=True, size=8)
        add_text(value_cell.paragraphs[0], value, size=8)
    doc.add_paragraph()


def add_field_card(cell: Any, field_item: SurveyField) -> None:
    cell.text = ""
    set_cell_margins(cell, top=140, start=140, bottom=140, end=140)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    label = cell.paragraphs[0]
    add_text(label, field_item.label, bold=True, size=7, color=TITLE_COLOR)
    label.paragraph_format.space_after = Pt(2)

    value = cell.add_paragraph()
    add_text(value, field_placeholder(field_item), size=9)
    value.paragraph_format.space_after = Pt(0)


def add_empty_card(cell: Any) -> None:
    cell.text = ""
    set_cell_margins(cell)
    set_shading(cell, "F7F9FA")


def add_field_card_table(doc: Document, fields: list[SurveyField]) -> None:
    if not fields:
        return

    row_count = (len(fields) + CARD_COLUMNS - 1) // CARD_COLUMNS
    table = doc.add_table(rows=row_count, cols=CARD_COLUMNS)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="E1E6EA", size="4")

    for row_index in range(row_count):
        for column_index in range(CARD_COLUMNS):
            field_index = row_index * CARD_COLUMNS + column_index
            cell = table.rows[row_index].cells[column_index]
            set_cell_width(cell, Inches(3.45))
            if field_index < len(fields):
                add_field_card(cell, fields[field_index])
            else:
                add_empty_card(cell)
    doc.add_paragraph()


def flatten_group_fields(group: SurveyGroup) -> list[tuple[str, SurveyField | None]]:
    rows: list[tuple[str, SurveyField | None]] = []
    if group.label:
        rows.append((group.label, None))
    for field_item in group.fields:
        rows.append((field_item.label, field_item))
    for child_group in group.groups:
        rows.extend(flatten_group_fields(child_group))
    for repeat in group.repeats:
        rows.append((repeat.label or repeat.name, None))
    return rows


def repeat_field_sections(repeat: SurveyRepeat) -> list[tuple[str, list[SurveyField]]]:
    sections: list[tuple[str, list[SurveyField]]] = []

    if repeat.fields:
        sections.append(("Details", repeat.fields))

    def collect_group_sections(group: SurveyGroup, prefix: str = "") -> None:
        label = group.label or group.name
        section_label = f"{prefix} - {label}" if prefix and label else label or prefix
        if group.fields:
            sections.append((section_label or "Details", group.fields))
        for child_group in group.groups:
            collect_group_sections(child_group, section_label)
        for child_repeat in group.repeats:
            for nested_label, nested_fields in repeat_field_sections(child_repeat):
                sections.append(
                    (
                        f"{section_label} - {child_repeat.label or child_repeat.name} - {nested_label}",
                        nested_fields,
                    )
                )

    for group_item in repeat.groups:
        collect_group_sections(group_item)

    return [(label, fields) for label, fields in sections if fields]


def repeat_rows(repeat: SurveyRepeat) -> list[tuple[str, SurveyField | None]]:
    rows: list[tuple[str, SurveyField | None]] = []
    for section_label, fields in repeat_field_sections(repeat):
        rows.append((section_label, None))
        rows.extend((field_item.label, field_item) for field_item in fields)
    return rows


def add_repeat_record_table(doc: Document, repeat: SurveyRepeat) -> None:
    rows = repeat_rows(repeat)
    if not rows:
        return

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="D8DEE3", size="5")

    field_header, response_header = table.rows[0].cells
    for cell in (field_header, response_header):
        cell.text = ""
        set_shading(cell, "EAF3F6")
        set_cell_margins(cell)
    set_cell_width(field_header, Inches(2.55))
    set_cell_width(response_header, Inches(4.45))
    add_text(field_header.paragraphs[0], "Field", bold=True, size=8, color=TITLE_COLOR)
    add_text(response_header.paragraphs[0], "Response", bold=True, size=8, color=TITLE_COLOR)

    for row_index, (label, field_item) in enumerate(rows, start=1):
        label_cell, value_cell = table.rows[row_index].cells
        label_cell.text = ""
        value_cell.text = ""
        set_cell_width(label_cell, Inches(2.55))
        set_cell_width(value_cell, Inches(4.45))
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        if field_item is None:
            merged_cell = label_cell.merge(value_cell)
            merged_cell.text = ""
            set_shading(merged_cell, "F0F4F6")
            set_cell_margins(merged_cell, top=90, bottom=90)
            add_text(merged_cell.paragraphs[0], label, bold=True, size=8, color=TITLE_COLOR)
            continue

        add_text(label_cell.paragraphs[0], label, bold=True, size=8)
        add_text(value_cell.paragraphs[0], field_placeholder(field_item), size=9)
    doc.add_paragraph()


def add_repeat_table(doc: Document, repeat: SurveyRepeat) -> None:
    rows = repeat_rows(repeat)
    start = doc.add_paragraph()
    add_text(start, f"${{# {repeat.name}}}", bold=True, size=8, color=TITLE_COLOR)
    add_section_band(doc, repeat.label or repeat.name, level=2)

    if not rows:
        paragraph = doc.add_paragraph()
        add_text(paragraph, f"${{{repeat.name}}}", size=8)
    else:
        add_repeat_record_table(doc, repeat)

    end = doc.add_paragraph()
    add_text(end, "${/}", bold=True, size=8, color=TITLE_COLOR)


def add_cover(doc: Document, form: SurveyForm) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(form.title)
    title_run.bold = True
    title_run.font.name = BODY_FONT
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = TITLE_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(subtitle, "ArcGIS Online Feature Report Template", bold=True, size=10)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(
        intro,
        "Designed for field review: compact field cards, readable response areas, and repeat records rendered as tables.",
        size=8,
    )

    meta_rows = [
        ("Survey Version", "${survey_version}"),
        ("Report Date", '${$date | format:"MM/DD/YYYY HH:mm"}'),
    ]
    if form.version:
        meta_rows.append(("XLSForm Version", form.version))
    add_label_value_table(doc, meta_rows, label_width=2.0)


def render_fields(doc: Document, fields: list[SurveyField]) -> None:
    add_field_card_table(doc, fields)


def render_group(doc: Document, group: SurveyGroup, depth: int = 1) -> None:
    if group.label:
        add_section_band(doc, group.label, level=1 if depth == 1 else 2)

    if group.fields:
        render_fields(doc, group.fields)

    for repeat in group.repeats:
        add_repeat_table(doc, repeat)

    for child_group in group.groups:
        render_group(doc, child_group, depth + 1)


def render_form(doc: Document, form: SurveyForm) -> None:
    for item in form.items:
        if isinstance(item, SurveyField):
            render_fields(doc, [item])
        elif isinstance(item, SurveyRepeat):
            add_repeat_table(doc, item)
        else:
            render_group(doc, item)


def generate_report_template(xlsx_path: Path, output_path: Path) -> Path:
    form = parse_xlsform(xlsx_path)
    document = Document()
    add_cover(document, form)
    render_form(document, form)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    resolved_output_path = next_available_path(output_path)
    document.save(resolved_output_path)
    return resolved_output_path


def next_available_path(path: Path) -> Path:
    """Return a non-existing path by appending a numeric suffix when needed."""
    if not path.exists():
        return path

    parent = path.parent
    stem = path.stem
    suffix = path.suffix
    counter = 1
    while True:
        candidate = parent / f"{stem}-{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a Survey123 feature report Word template from an XLSForm."
    )
    parser.add_argument(
        "--xlsx",
        required=True,
        type=Path,
        help="Path to the Survey123 Connect XLSForm (.xlsx).",
    )
    parser.add_argument(
        "--output",
        required=True,
        type=Path,
        help="Output .docx path for the generated ArcGIS Online feature report template.",
    )
    args = parser.parse_args()

    if not args.xlsx.exists():
        raise SystemExit(f"XLSForm not found: {args.xlsx}")

    output = generate_report_template(args.xlsx, args.output)
    print(f"Wrote report template: {output}")


if __name__ == "__main__":
    main()
