#!/usr/bin/env python3
"""Generate an ArcGIS Survey123 feature report Word template from an XLSForm.

Bootstraps page layout and styles from a reference report template (.docx),
then builds sections, label/value tables, yes/no checkboxes, and repeat grids
from the Survey123 Connect workbook.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from html import unescape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook

BODY_FONT = "Arial"
TITLE_COLOR = RGBColor(0x0A, 0x1F, 0x62)
SECTION_FILL = "0A1F62"
SUBSECTION_FILL = "D9EAF0"
LABEL_FILL = "EAF3F6"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
REPEAT_COLUMNS_PER_TABLE = 5
PAGE_WIDTH_INCHES = 7.0

SKIP_TYPES = {
    "note",
    "calculate",
    "begin group",
    "end group",
    "begin repeat",
    "end repeat",
}

MEDIA_TYPES = {"image", "audio", "file"}
MAP_TYPES = {"geopoint", "geotrace", "geoshape"}
DATE_TYPES = {"date", "datetime"}


@dataclass
class Choice:
    name: str
    label: str


@dataclass
class SurveyField:
    type: str
    name: str
    label: str
    appearance: str = ""
    list_name: str | None = None
    choices: list[Choice] = field(default_factory=list)


@dataclass
class SurveyGroup:
    name: str
    label: str
    fields: list[SurveyField] = field(default_factory=list)
    groups: list["SurveyGroup"] = field(default_factory=list)
    repeats: list["SurveyRepeat"] = field(default_factory=list)


@dataclass
class SurveyRepeat:
    name: str
    label: str
    fields: list[SurveyField] = field(default_factory=list)
    groups: list[SurveyGroup] = field(default_factory=list)


@dataclass
class SurveyForm:
    title: str
    version: str
    items: list[SurveyField | SurveyGroup | SurveyRepeat] = field(default_factory=list)


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = unescape(str(value))
    text = re.sub(r"<br\s*/?>", " ", text, flags=re.I)
    text = re.sub(r"</(?:p|h\d|div|li)>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("\xa0", " ")
    return re.sub(r"\s+", " ", text).strip()


def base_type(type_name: str) -> str:
    return type_name.strip().split(maxsplit=1)[0].lower()


def choice_list_from_type(type_name: str) -> str | None:
    match = re.match(
        r"select_(?:one|multiple)(?:_from_file)?\s+(.+)",
        type_name.strip(),
        flags=re.I,
    )
    return match.group(1).strip() if match else None


def load_choices(workbook) -> dict[str, list[Choice]]:
    if "choices" not in workbook.sheetnames:
        return {}

    worksheet = workbook["choices"]
    rows = worksheet.iter_rows(values_only=True)
    header = next(rows, None)
    if not header:
        return {}

    columns = {str(name).strip(): i for i, name in enumerate(header) if name}
    list_idx = columns.get("list_name", 0)
    name_idx = columns.get("name", 1)
    label_idx = columns.get("label", 2)

    choices: dict[str, list[Choice]] = {}
    for row in rows:
        if not row or list_idx >= len(row) or not row[list_idx]:
            continue
        list_name = str(row[list_idx]).strip()
        choice_name = row[name_idx] if name_idx < len(row) else None
        if choice_name is None or str(choice_name).strip() == "":
            continue
        choice_label = row[label_idx] if label_idx < len(row) else choice_name
        choices.setdefault(list_name, []).append(
            Choice(str(choice_name).strip(), clean_text(choice_label) or str(choice_name))
        )
    return choices


def load_settings(workbook) -> tuple[str, str]:
    if "settings" not in workbook.sheetnames:
        return "Survey Feature Report", ""

    worksheet = workbook["settings"]
    rows = list(worksheet.iter_rows(values_only=True))
    if len(rows) < 2:
        return "Survey Feature Report", ""

    header = [str(value).strip() if value else "" for value in rows[0]]
    values = dict(zip(header, rows[1]))
    title = clean_text(values.get("form_title")) or "Survey Feature Report"
    version = str(values.get("version") or "").strip()
    return title, version


def row_value(row: tuple[Any, ...], columns: dict[str, int], column_name: str) -> Any:
    index = columns.get(column_name)
    if index is None or index >= len(row):
        return None
    return row[index]


def should_skip_field(row: dict[str, Any]) -> bool:
    field_type = str(row.get("type") or "").strip().lower()
    field_name = str(row.get("name") or "").strip()
    if not field_name or field_name == "list_name":
        return True
    if field_type in SKIP_TYPES or base_type(field_type) in SKIP_TYPES:
        return True
    if str(row.get("visible") or "").strip().lower() == "no":
        return True
    if str(row.get("field_type") or "").strip().lower() == "null":
        return True
    return False


def make_field(row: dict[str, Any], choices: dict[str, list[Choice]]) -> SurveyField:
    type_name = str(row["type"]).strip()
    list_name = choice_list_from_type(type_name)
    label = clean_text(row.get("label")) or clean_text(row.get("alias"))
    return SurveyField(
        type=type_name,
        name=str(row["name"]).strip(),
        label=label or str(row["name"]).strip(),
        appearance=str(row.get("appearance") or "").strip(),
        list_name=list_name,
        choices=choices.get(list_name or "", []),
    )


def parse_xlsform(xlsx_path: Path) -> SurveyForm:
    workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
    try:
        if "survey" not in workbook.sheetnames:
            raise ValueError(f"Workbook does not contain a survey sheet: {xlsx_path}")

        title, version = load_settings(workbook)
        choices = load_choices(workbook)
        worksheet = workbook["survey"]

        header_row = next(worksheet.iter_rows(max_row=1, values_only=True))
        header = [str(value).strip() if value else "" for value in header_row]
        columns = {name: index for index, name in enumerate(header) if name}

        root_items: list[SurveyField | SurveyGroup | SurveyRepeat] = []
        stack: list[SurveyGroup | SurveyRepeat] = []

        def current_parent() -> SurveyGroup | SurveyRepeat | None:
            return stack[-1] if stack else None

        def append_item(item: SurveyField | SurveyGroup | SurveyRepeat) -> None:
            parent = current_parent()
            if parent is None:
                root_items.append(item)
            elif isinstance(parent, SurveyGroup):
                if isinstance(item, SurveyField):
                    parent.fields.append(item)
                elif isinstance(item, SurveyRepeat):
                    parent.repeats.append(item)
                else:
                    parent.groups.append(item)
            elif isinstance(item, SurveyField):
                parent.fields.append(item)
            elif isinstance(item, SurveyGroup):
                parent.groups.append(item)
            else:
                root_items.append(item)

        for row in worksheet.iter_rows(min_row=2, values_only=True):
            mapped = {
                "type": row_value(row, columns, "type"),
                "name": row_value(row, columns, "name"),
                "label": row_value(row, columns, "label"),
                "appearance": row_value(row, columns, "appearance"),
                "visible": row_value(row, columns, "body::esri:visible"),
                "field_type": row_value(row, columns, "bind::esri:fieldType"),
                "alias": row_value(row, columns, "bind::esri:fieldAlias"),
            }
            row_type = str(mapped["type"] or "").strip()
            normalized_type = row_type.lower()
            if not row_type:
                continue

            if normalized_type == "begin group":
                group = SurveyGroup(
                    name=str(mapped.get("name") or "").strip(),
                    label=clean_text(mapped.get("label"))
                    or str(mapped.get("name") or "").strip(),
                )
                append_item(group)
                stack.append(group)
                continue

            if normalized_type == "end group":
                while stack and not isinstance(stack[-1], SurveyGroup):
                    stack.pop()
                if stack:
                    stack.pop()
                continue

            if normalized_type == "begin repeat":
                repeat = SurveyRepeat(
                    name=str(mapped.get("name") or "").strip(),
                    label=clean_text(mapped.get("label"))
                    or str(mapped.get("name") or "").strip(),
                )
                append_item(repeat)
                stack.append(repeat)
                continue

            if normalized_type == "end repeat":
                while stack and not isinstance(stack[-1], SurveyRepeat):
                    stack.pop()
                if stack:
                    stack.pop()
                continue

            if not should_skip_field(mapped):
                append_item(make_field(mapped, choices))

        return SurveyForm(title=title, version=version, items=root_items)
    finally:
        workbook.close()


def clear_document_body(document: Document) -> None:
    body = document.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            body.remove(child)


def choice_list_name(field_item: SurveyField) -> str | None:
    return field_item.list_name or choice_list_from_type(field_item.type)


def is_yes_only(field_item: SurveyField) -> bool:
    list_name = choice_list_name(field_item)
    if list_name == "yes_only":
        return True
    choice_names = {choice.name.lower() for choice in field_item.choices}
    return choice_names == {"yes"}


def is_yes_no_family(field_item: SurveyField) -> bool:
    list_name = choice_list_name(field_item)
    return list_name in {"yes_no", "yes_no_unk", "yes_no_na"}


def checkbox(name: str, value: str, label: str | None = None) -> str:
    text = f'${{{name} | checked:"{value}"}}'
    return f"{text} {label}" if label else text


def field_placeholder(field_item: SurveyField, include_choice_labels: bool = True) -> str:
    type_name = base_type(field_item.type)
    name = field_item.name

    if is_yes_only(field_item):
        return checkbox(name, "yes")

    if is_yes_no_family(field_item):
        parts = [checkbox(name, "yes", "Yes"), checkbox(name, "no", "No")]
        list_name = choice_list_name(field_item)
        if list_name == "yes_no_unk":
            parts.append(checkbox(name, "unknown", "Unknown"))
        elif list_name == "yes_no_na":
            parts.append(checkbox(name, "n/a", "N/A"))
        return "  ".join(parts)

    if type_name.startswith("select_one") and field_item.choices and len(field_item.choices) <= 10:
        return "  ".join(
            checkbox(name, choice.name, choice.label if include_choice_labels else None)
            for choice in field_item.choices
        )

    if type_name.startswith("select_multiple"):
        if field_item.choices and len(field_item.choices) <= 10:
            return "  ".join(
                checkbox(name, choice.name, choice.label if include_choice_labels else None)
                for choice in field_item.choices
            )
        return f'${{{name} | appearance:"bullets"}}'

    if type_name in DATE_TYPES:
        return f'${{{name} | format:"MM/DD/YYYY"}}'

    if type_name in MEDIA_TYPES:
        return (
            f"${{if {name}}}"
            f"${{# {name}}}"
            f'${{$file | getValue:"" | size:0:0:600:450}}'
            f"${{/}}"
            f"${{/}}"
        )

    if type_name in MAP_TYPES:
        return f"${{{name}}}"

    if type_name == "text" and "multiline" in field_item.appearance.lower():
        return f'${{{name} | appearance:"multiline"}}'

    return f"${{{name}}}"


def set_shading(cell_or_paragraph: Any, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    if hasattr(cell_or_paragraph, "_tc"):
        cell_or_paragraph._tc.get_or_add_tcPr().append(shading)
    else:
        cell_or_paragraph._p.get_or_add_pPr().append(shading)


def set_cell_width(cell: Any, width: Inches) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_text(paragraph: Any, text: str, *, bold: bool = False, size: int = 9) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT
    run.font.size = Pt(size)


def add_section_heading(document: Document, text: str, level: int = 1) -> None:
    if not text:
        return
    if level == 1:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_shading(paragraph, SECTION_FILL)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(12)
        run.font.color.rgb = HEADER_TEXT
        return

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = TITLE_COLOR


def add_label_value_table(
    document: Document, rows: list[tuple[str, str]], label_width: float = 2.6
) -> None:
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[index].cells
        label_cell.text = ""
        value_cell.text = ""
        set_cell_width(label_cell, Inches(label_width))
        set_cell_width(value_cell, Inches(PAGE_WIDTH_INCHES - label_width))
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_shading(label_cell, LABEL_FILL)
        add_text(label_cell.paragraphs[0], label, bold=True, size=8)
        add_text(value_cell.paragraphs[0], value, size=8)
    document.add_paragraph()


def chunked(items: list[SurveyField], size: int) -> list[list[SurveyField]]:
    return [items[index : index + size] for index in range(0, len(items), size)]


def collect_repeat_fields(repeat: SurveyRepeat) -> list[SurveyField]:
    fields = list(repeat.fields)

    def walk_group(group: SurveyGroup) -> None:
        fields.extend(group.fields)
        for child in group.groups:
            walk_group(child)

    for group_item in repeat.groups:
        walk_group(group_item)

    return fields


def add_repeat_table(document: Document, repeat: SurveyRepeat) -> None:
    add_section_heading(document, repeat.label or repeat.name, level=2)
    fields = collect_repeat_fields(repeat)
    if not fields:
        paragraph = document.add_paragraph()
        add_text(paragraph, f"${{# {repeat.name}}}${{{repeat.name}}}${{/}}", size=8)
        return

    for field_chunk in chunked(fields, REPEAT_COLUMNS_PER_TABLE):
        table = document.add_table(rows=2, cols=len(field_chunk))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        column_width = PAGE_WIDTH_INCHES / max(len(field_chunk), 1)

        for index, field_item in enumerate(field_chunk):
            header_cell = table.rows[0].cells[index]
            value_cell = table.rows[1].cells[index]
            header_cell.text = ""
            value_cell.text = ""
            set_shading(header_cell, LABEL_FILL)
            header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(header_cell, Inches(column_width))
            set_cell_width(value_cell, Inches(column_width))

            add_text(header_cell.paragraphs[0], field_item.label, bold=True, size=7)
            placeholder = field_placeholder(field_item)
            if index == 0:
                placeholder = f"${{# {repeat.name}}}{placeholder}"
            if index == len(field_chunk) - 1:
                placeholder = f"{placeholder}${{/}}"
            add_text(value_cell.paragraphs[0], placeholder, size=7)

        document.add_paragraph()


def add_cover_from_template(document: Document, form: SurveyForm) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(form.title)
    title_run.bold = True
    title_run.font.name = BODY_FONT
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = TITLE_COLOR

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("SITE SURVEY WORKBOOK")
    subtitle_run.bold = True
    subtitle_run.font.name = BODY_FONT
    subtitle_run.font.size = Pt(26)
    subtitle_run.font.color.rgb = TITLE_COLOR

    document.add_paragraph()

    version_line = document.add_paragraph()
    version_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_text = "Report Version 1.0\nSurvey Version ${survey_version}"
    if form.version:
        version_text += f"\nXLSForm Version {form.version}"
    version_run = version_line.add_run(version_text)
    version_run.font.name = BODY_FONT
    version_run.font.size = Pt(12)

    document.add_paragraph()
    document.add_paragraph()

    date_heading = document.add_paragraph()
    date_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_heading.add_run("(Date)")
    date_run.bold = True
    date_run.font.name = BODY_FONT
    date_run.font.size = Pt(16)

    date_table = document.add_table(rows=1, cols=3)
    date_table.style = "Table Grid"
    date_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    center = date_table.rows[0].cells[1]
    center.text = ""
    add_text(
        center.paragraphs[0],
        '${survey_date | format:"MM/DD/YYYY HH:mm"}',
        bold=False,
        size=10,
    )
    document.add_paragraph()


def render_fields(document: Document, fields: list[SurveyField]) -> None:
    rows = [(field_item.label, field_placeholder(field_item)) for field_item in fields]
    add_label_value_table(document, rows)


def render_group(document: Document, group: SurveyGroup, depth: int = 1) -> None:
    if group.label:
        add_section_heading(document, group.label, level=1 if depth == 1 else 2)

    if group.fields:
        render_fields(document, group.fields)

    for repeat in group.repeats:
        add_repeat_table(document, repeat)

    for child_group in group.groups:
        render_group(document, child_group, depth + 1)


def render_form(document: Document, form: SurveyForm) -> None:
    for item in form.items:
        if isinstance(item, SurveyField):
            render_fields(document, [item])
        elif isinstance(item, SurveyRepeat):
            add_repeat_table(document, item)
        else:
            render_group(document, item)


def load_document_from_template(template_path: Path) -> Document:
    document = Document(str(template_path))
    clear_document_body(document)
    return document


def generate_report_template(
    xlsx_path: Path, template_path: Path, output_path: Path
) -> Path:
    form = parse_xlsform(xlsx_path)
    document = load_document_from_template(template_path)
    add_cover_from_template(document, form)
    render_form(document, form)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Generate a Survey123 ArcGIS Online feature report Word template "
            "from an XLSForm, using a reference .docx for layout."
        )
    )
    parser.add_argument(
        "--xlsx",
        required=True,
        type=Path,
        help="Path to the Survey123 Connect XLSForm (.xlsx).",
    )
    parser.add_argument(
        "--template",
        required=True,
        type=Path,
        help="Reference feature report .docx (layout mockup; body content is replaced).",
    )
    parser.add_argument(
        "--output",
        required=True,
        type=Path,
        help="Output .docx path for the generated feature report template.",
    )
    args = parser.parse_args()

    if not args.xlsx.exists():
        raise SystemExit(f"XLSForm not found: {args.xlsx}")
    if not args.template.exists():
        raise SystemExit(f"Template not found: {args.template}")

    output = generate_report_template(args.xlsx, args.template, args.output)
    print(f"Wrote report template: {output}")


if __name__ == "__main__":
    main()
