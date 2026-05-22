#!/usr/bin/env python3
"""LPGBS Inspection Dossier - Survey123 feature report template generator.

Reads a Survey123 Connect XLSForm and produces a Word (.docx) ArcGIS Online
feature report template with a distinctive "inspection dossier" look:

    * Auto-numbered section / subsection banners (1.0, 1.1, ...)
    * Amber sidebar accent on every subsection block
    * Compact 3-column "checkbox grid" for short questions
      (yes_only, yes_no, integer, decimal, date)
    * Full-width "response card" for long-text / multi-choice / media fields
    * Vertical, iterating two-column tables for repeats
    * `note` field types are skipped per requirement
    * Yes-only choices render as a single ArcGIS Survey123 checkbox glyph

The Survey123 checkbox syntax
(`${field | checked:"choice_name"}`) renders a filled or empty box at
report-generation time, so every "checkbox" in the generated template is
both a visual marker and a working placeholder.

Usage:
    python tools/lpgbs_inspection_report_generator.py \
        --xlsx "C:/path/to/LPGBS Site Survey Streamlined Only.xlsx" \
        --output outputs/lpgbs-inspection-report.docx
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from html import unescape
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

NAVY = RGBColor(0x0A, 0x1F, 0x62)
AMBER = RGBColor(0xC8, 0x96, 0x1F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SLATE_DARK = RGBColor(0x33, 0x39, 0x44)
SLATE_MID = RGBColor(0x5A, 0x60, 0x6E)

NAVY_HEX = "0A1F62"
AMBER_HEX = "C8961F"
CREAM_HEX = "FAF7EE"
SUBSECTION_HEX = "EAEEF6"
DIVIDER_HEX = "C5CBD8"
RESPONSE_HEX = "FFFFFF"
SIDEBAR_HEX = AMBER_HEX

PAGE_WIDTH_IN = 7.1
GRID_COLUMNS = 3
SIDEBAR_WIDTH_IN = 0.08

SKIP_TYPES = {"note", "calculate"}
SHORT_BASE_TYPES = {"integer", "decimal", "date", "time", "datetime"}
MEDIA_TYPES = {"image", "audio", "file"}
MAP_TYPES = {"geopoint", "geotrace", "geoshape"}


@dataclass
class Choice:
    name: str
    label: str


@dataclass
class SurveyField:
    type: str
    name: str
    label: str
    appearance: str = ""
    list_name: str | None = None
    choices: list[Choice] = field(default_factory=list)


@dataclass
class SurveyGroup:
    name: str
    label: str
    fields: list[SurveyField] = field(default_factory=list)
    groups: list["SurveyGroup"] = field(default_factory=list)
    repeats: list["SurveyRepeat"] = field(default_factory=list)


@dataclass
class SurveyRepeat:
    name: str
    label: str
    fields: list[SurveyField] = field(default_factory=list)
    groups: list[SurveyGroup] = field(default_factory=list)


@dataclass
class SurveyForm:
    title: str
    form_id: str
    version: str
    items: list[SurveyField | SurveyGroup | SurveyRepeat] = field(default_factory=list)


# ---------------------------------------------------------------------------
# XLSForm parsing
# ---------------------------------------------------------------------------


def clean_text(value: Any) -> str:
    """Return plain text from an XLSForm label / hint that may contain HTML."""
    if value is None:
        return ""
    text = unescape(str(value))
    text = re.sub(r"<br\s*/?>", " ", text, flags=re.I)
    text = re.sub(r"</(?:p|h\d|div|li|tr|td)>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def base_type(type_name: str) -> str:
    return type_name.strip().split(maxsplit=1)[0].lower()


def choice_list_from_type(type_name: str) -> str | None:
    match = re.match(
        r"select_(?:one|multiple)(?:_from_file)?\s+(.+)",
        type_name.strip(),
        flags=re.I,
    )
    return match.group(1).strip() if match else None


def load_choices(workbook) -> dict[str, list[Choice]]:
    if "choices" not in workbook.sheetnames:
        return {}

    worksheet = workbook["choices"]
    rows = worksheet.iter_rows(values_only=True)
    header = next(rows, None)
    if not header:
        return {}

    columns = {str(name).strip(): index for index, name in enumerate(header) if name}
    list_idx = columns.get("list_name", 0)
    name_idx = columns.get("name", 1)
    label_idx = columns.get("label", 2)

    choices: dict[str, list[Choice]] = {}
    for row in rows:
        if not row or list_idx >= len(row) or not row[list_idx]:
            continue
        list_name = str(row[list_idx]).strip()
        raw_name = row[name_idx] if name_idx < len(row) else None
        if raw_name is None or str(raw_name).strip() == "":
            continue
        raw_label = row[label_idx] if label_idx < len(row) else raw_name
        choices.setdefault(list_name, []).append(
            Choice(
                name=str(raw_name).strip(),
                label=clean_text(raw_label) or str(raw_name).strip(),
            )
        )
    return choices


def load_settings(workbook) -> tuple[str, str, str]:
    if "settings" not in workbook.sheetnames:
        return "Survey Feature Report", "", ""

    worksheet = workbook["settings"]
    rows = list(worksheet.iter_rows(values_only=True))
    if len(rows) < 2:
        return "Survey Feature Report", "", ""

    header = [str(value).strip() if value else "" for value in rows[0]]
    values = dict(zip(header, rows[1]))
    title = clean_text(values.get("form_title")) or "Survey Feature Report"
    form_id = str(values.get("form_id") or "").strip()
    version = str(values.get("version") or "").strip()
    return title, form_id, version


def row_value(row: tuple[Any, ...], columns: dict[str, int], column_name: str) -> Any:
    index = columns.get(column_name)
    if index is None or index >= len(row):
        return None
    return row[index]


def should_skip_field(mapped: dict[str, Any]) -> bool:
    field_type = str(mapped.get("type") or "").strip().lower()
    field_name = str(mapped.get("name") or "").strip()
    if not field_name:
        return True
    if field_type in SKIP_TYPES or base_type(field_type) in SKIP_TYPES:
        return True
    if str(mapped.get("visible") or "").strip().lower() == "no":
        return True
    if str(mapped.get("field_type") or "").strip().lower() == "null":
        return True
    return False


def make_field(mapped: dict[str, Any], choices: dict[str, list[Choice]]) -> SurveyField:
    type_name = str(mapped["type"]).strip()
    list_name = choice_list_from_type(type_name)
    label = clean_text(mapped.get("label")) or clean_text(mapped.get("alias"))
    return SurveyField(
        type=type_name,
        name=str(mapped["name"]).strip(),
        label=label or str(mapped["name"]).strip(),
        appearance=str(mapped.get("appearance") or "").strip(),
        list_name=list_name,
        choices=choices.get(list_name or "", []),
    )


def parse_xlsform(xlsx_path: Path) -> SurveyForm:
    workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
    try:
        if "survey" not in workbook.sheetnames:
            raise ValueError(f"Workbook is missing the 'survey' sheet: {xlsx_path}")

        title, form_id, version = load_settings(workbook)
        choices = load_choices(workbook)
        worksheet = workbook["survey"]

        header_row = next(worksheet.iter_rows(max_row=1, values_only=True))
        header = [str(value).strip() if value else "" for value in header_row]
        columns = {name: index for index, name in enumerate(header) if name}

        root_items: list[SurveyField | SurveyGroup | SurveyRepeat] = []
        stack: list[SurveyGroup | SurveyRepeat] = []

        def parent() -> SurveyGroup | SurveyRepeat | None:
            return stack[-1] if stack else None

        def append(item: SurveyField | SurveyGroup | SurveyRepeat) -> None:
            head = parent()
            if head is None:
                root_items.append(item)
            elif isinstance(head, SurveyGroup):
                if isinstance(item, SurveyField):
                    head.fields.append(item)
                elif isinstance(item, SurveyRepeat):
                    head.repeats.append(item)
                else:
                    head.groups.append(item)
            else:
                if isinstance(item, SurveyField):
                    head.fields.append(item)
                elif isinstance(item, SurveyGroup):
                    head.groups.append(item)
                else:
                    root_items.append(item)

        for row in worksheet.iter_rows(min_row=2, values_only=True):
            mapped = {
                "type": row_value(row, columns, "type"),
                "name": row_value(row, columns, "name"),
                "label": row_value(row, columns, "label"),
                "appearance": row_value(row, columns, "appearance"),
                "visible": row_value(row, columns, "body::esri:visible"),
                "field_type": row_value(row, columns, "bind::esri:fieldType"),
                "alias": row_value(row, columns, "bind::esri:fieldAlias"),
            }
            raw_type = str(mapped["type"] or "").strip()
            kind = raw_type.lower()
            if not raw_type:
                continue

            if kind == "begin group":
                group = SurveyGroup(
                    name=str(mapped.get("name") or "").strip(),
                    label=clean_text(mapped.get("label"))
                    or str(mapped.get("name") or "").strip(),
                )
                append(group)
                stack.append(group)
                continue

            if kind == "end group":
                while stack and not isinstance(stack[-1], SurveyGroup):
                    stack.pop()
                if stack:
                    stack.pop()
                continue

            if kind == "begin repeat":
                repeat = SurveyRepeat(
                    name=str(mapped.get("name") or "").strip(),
                    label=clean_text(mapped.get("label"))
                    or str(mapped.get("name") or "").strip(),
                )
                append(repeat)
                stack.append(repeat)
                continue

            if kind == "end repeat":
                while stack and not isinstance(stack[-1], SurveyRepeat):
                    stack.pop()
                if stack:
                    stack.pop()
                continue

            if not should_skip_field(mapped):
                append(make_field(mapped, choices))

        return SurveyForm(title=title, form_id=form_id, version=version, items=root_items)
    finally:
        workbook.close()


# ---------------------------------------------------------------------------
# Field classification
# ---------------------------------------------------------------------------


def choice_list_name(field_item: SurveyField) -> str | None:
    return field_item.list_name or choice_list_from_type(field_item.type)


def is_yes_only(field_item: SurveyField) -> bool:
    if choice_list_name(field_item) == "yes_only":
        return True
    names = {choice.name.strip().lower() for choice in field_item.choices}
    return names == {"yes"}


def is_yes_no_family(field_item: SurveyField) -> bool:
    return choice_list_name(field_item) in {"yes_no", "yes_no_unk", "yes_no_na"}


def is_short_field(field_item: SurveyField) -> bool:
    """True if a field can live comfortably inside a 1/3-page grid cell."""
    if is_yes_only(field_item) or is_yes_no_family(field_item):
        return True
    type_name = base_type(field_item.type)
    if type_name in SHORT_BASE_TYPES:
        return True
    if type_name.startswith("select_one") and field_item.choices:
        if "multiline" in field_item.appearance.lower():
            return False
        return len(field_item.choices) <= 3
    if type_name == "text":
        return "multiline" not in field_item.appearance.lower()
    return False


def is_full_width_field(field_item: SurveyField) -> bool:
    type_name = base_type(field_item.type)
    if type_name in MEDIA_TYPES or type_name in MAP_TYPES:
        return True
    if type_name.startswith("select_multiple"):
        return True
    if type_name.startswith("select_one") and not is_short_field(field_item):
        return True
    if type_name == "text" and "multiline" in field_item.appearance.lower():
        return True
    return False


# ---------------------------------------------------------------------------
# Survey123 placeholders
# ---------------------------------------------------------------------------


def checkbox_placeholder(name: str, value: str) -> str:
    return f'${{{name} | checked:"{value}"}}'


def labeled_checkbox(name: str, value: str, label: str | None) -> str:
    glyph = checkbox_placeholder(name, value)
    return f"{glyph} {label}" if label else glyph


def field_response(field_item: SurveyField) -> str:
    """Survey123 placeholder string for a non-checkbox field response."""
    name = field_item.name
    type_name = base_type(field_item.type)

    if type_name in {"date"}:
        return f'${{{name} | format:"MM/DD/YYYY"}}'
    if type_name in {"datetime"}:
        return f'${{{name} | format:"MM/DD/YYYY hh:mm a"}}'
    if type_name == "time":
        return f'${{{name} | format:"hh:mm a"}}'
    if type_name in MEDIA_TYPES:
        return (
            f"${{# {name}}}"
            f'${{$file | size:0:0:560:380}}'
            f"${{/}}"
        )
    if type_name in MAP_TYPES:
        return f"${{{name}}}"
    if type_name.startswith("select_multiple"):
        return f'${{{name} | appearance:"bullets"}}'
    if type_name == "text" and "multiline" in field_item.appearance.lower():
        return f'${{{name} | appearance:"multiline"}}'
    return f"${{{name}}}"


# ---------------------------------------------------------------------------
# OXML helpers
# ---------------------------------------------------------------------------


def _shading_element(fill_hex: str) -> OxmlElement:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    return shading


def set_cell_shading(cell: Any, fill_hex: str) -> None:
    cell._tc.get_or_add_tcPr().append(_shading_element(fill_hex))


def set_paragraph_shading(paragraph: Any, fill_hex: str) -> None:
    paragraph._p.get_or_add_pPr().append(_shading_element(fill_hex))


def set_cell_width(cell: Any, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(
    cell: Any,
    *,
    top: int = 80,
    start: int = 110,
    bottom: int = 80,
    end: int = 110,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(
    table: Any,
    *,
    color: str = DIVIDER_HEX,
    size: str = "4",
    style: str = "single",
) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), style)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def remove_table_borders(table: Any) -> None:
    set_table_borders(table, color="FFFFFF", size="0", style="nil")


def add_text(
    paragraph: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 9,
    color: RGBColor | None = None,
    font: str = BODY_FONT,
) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT if font == BODY_FONT else font
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def clear_paragraph_spacing(paragraph: Any) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Visual building blocks
# ---------------------------------------------------------------------------


def add_section_banner(doc: Document, number: str, title: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    tag_cell, title_cell = table.rows[0].cells
    set_cell_width(tag_cell, 0.9)
    set_cell_width(title_cell, PAGE_WIDTH_IN - 0.9)
    set_cell_shading(tag_cell, AMBER_HEX)
    set_cell_shading(title_cell, NAVY_HEX)
    set_cell_margins(tag_cell, top=120, bottom=120, start=140, end=140)
    set_cell_margins(title_cell, top=120, bottom=120, start=180, end=180)
    tag_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tag_paragraph = tag_cell.paragraphs[0]
    tag_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph_spacing(tag_paragraph)
    add_text(tag_paragraph, number, bold=True, size=13, color=NAVY)

    title_paragraph = title_cell.paragraphs[0]
    clear_paragraph_spacing(title_paragraph)
    add_text(title_paragraph, title.upper(), bold=True, size=13, color=WHITE)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_subsection_banner(doc: Document, number: str, title: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    tag_cell, title_cell = table.rows[0].cells
    set_cell_width(tag_cell, 0.7)
    set_cell_width(title_cell, PAGE_WIDTH_IN - 0.7)
    set_cell_shading(tag_cell, NAVY_HEX)
    set_cell_shading(title_cell, SUBSECTION_HEX)
    set_cell_margins(tag_cell, top=70, bottom=70, start=100, end=100)
    set_cell_margins(title_cell, top=70, bottom=70, start=140, end=140)
    tag_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tag_paragraph = tag_cell.paragraphs[0]
    tag_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph_spacing(tag_paragraph)
    add_text(tag_paragraph, number, bold=True, size=10, color=WHITE)

    title_paragraph = title_cell.paragraphs[0]
    clear_paragraph_spacing(title_paragraph)
    add_text(title_paragraph, title, bold=True, size=10, color=NAVY)


def open_sidebar_frame(doc: Document) -> Any:
    """Return the content cell of an amber-sidebar wrapper table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    sidebar_cell, content_cell = table.rows[0].cells
    set_cell_width(sidebar_cell, SIDEBAR_WIDTH_IN)
    set_cell_width(content_cell, PAGE_WIDTH_IN - SIDEBAR_WIDTH_IN)
    set_cell_shading(sidebar_cell, SIDEBAR_HEX)
    set_cell_margins(sidebar_cell, top=20, bottom=20, start=0, end=0)
    set_cell_margins(content_cell, top=80, bottom=80, start=160, end=80)
    sidebar_paragraph = sidebar_cell.paragraphs[0]
    clear_paragraph_spacing(sidebar_paragraph)
    return content_cell


def short_field_block_text(field_item: SurveyField) -> tuple[str, str]:
    """Return (label, response_text) for a short-field grid cell."""
    label = field_item.label
    if is_yes_only(field_item):
        return label, labeled_checkbox(field_item.name, "yes", "Yes")
    if is_yes_no_family(field_item):
        parts = [
            labeled_checkbox(field_item.name, "yes", "Yes"),
            labeled_checkbox(field_item.name, "no", "No"),
        ]
        list_name = choice_list_name(field_item)
        if list_name == "yes_no_unk":
            parts.append(labeled_checkbox(field_item.name, "unknown", "Unknown"))
        elif list_name == "yes_no_na":
            parts.append(labeled_checkbox(field_item.name, "n/a", "N/A"))
        return label, "   ".join(parts)
    type_name = base_type(field_item.type)
    if type_name.startswith("select_one") and field_item.choices:
        glyphs = [
            labeled_checkbox(field_item.name, choice.name, choice.label)
            for choice in field_item.choices
        ]
        return label, "   ".join(glyphs)
    return label, field_response(field_item)


def render_short_field_grid(parent_cell: Any, fields: list[SurveyField]) -> None:
    if not fields:
        return
    rows_needed = (len(fields) + GRID_COLUMNS - 1) // GRID_COLUMNS
    grid = parent_cell.add_table(rows=rows_needed, cols=GRID_COLUMNS)
    grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(grid, color=DIVIDER_HEX, size="4")

    column_width = (PAGE_WIDTH_IN - SIDEBAR_WIDTH_IN - 0.25) / GRID_COLUMNS

    for index in range(rows_needed * GRID_COLUMNS):
        row_index, col_index = divmod(index, GRID_COLUMNS)
        cell = grid.rows[row_index].cells[col_index]
        cell.text = ""
        set_cell_width(cell, column_width)
        set_cell_margins(cell, top=80, bottom=80, start=110, end=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        if index >= len(fields):
            set_cell_shading(cell, CREAM_HEX)
            continue
        field_item = fields[index]
        label_text, response_text = short_field_block_text(field_item)
        label_paragraph = cell.paragraphs[0]
        clear_paragraph_spacing(label_paragraph)
        add_text(label_paragraph, label_text, bold=True, size=7, color=SLATE_MID)
        response_paragraph = cell.add_paragraph()
        clear_paragraph_spacing(response_paragraph)
        response_paragraph.paragraph_format.space_before = Pt(1)
        add_text(response_paragraph, response_text, size=9, color=SLATE_DARK)


def render_full_width_field(parent_cell: Any, field_item: SurveyField) -> None:
    table = parent_cell.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color=DIVIDER_HEX, size="4")
    width = PAGE_WIDTH_IN - SIDEBAR_WIDTH_IN - 0.25

    header = table.rows[0].cells[0]
    body = table.rows[1].cells[0]
    set_cell_width(header, width)
    set_cell_width(body, width)
    set_cell_shading(header, NAVY_HEX)
    set_cell_shading(body, RESPONSE_HEX)
    set_cell_margins(header, top=60, bottom=60, start=140, end=140)
    set_cell_margins(body, top=120, bottom=120, start=140, end=140)

    header_paragraph = header.paragraphs[0]
    clear_paragraph_spacing(header_paragraph)
    add_text(header_paragraph, field_item.label, bold=True, size=8, color=WHITE)

    body_paragraph = body.paragraphs[0]
    clear_paragraph_spacing(body_paragraph)

    type_name = base_type(field_item.type)
    if type_name.startswith("select_multiple") and field_item.choices and len(field_item.choices) <= 18:
        for index, choice in enumerate(field_item.choices):
            line = body_paragraph if index == 0 else body.add_paragraph()
            clear_paragraph_spacing(line)
            add_text(line, labeled_checkbox(field_item.name, choice.name, choice.label), size=9, color=SLATE_DARK)
        return

    if type_name.startswith("select_one") and field_item.choices and not is_short_field(field_item):
        for index, choice in enumerate(field_item.choices):
            line = body_paragraph if index == 0 else body.add_paragraph()
            clear_paragraph_spacing(line)
            add_text(line, labeled_checkbox(field_item.name, choice.name, choice.label), size=9, color=SLATE_DARK)
        return

    add_text(body_paragraph, field_response(field_item), size=9, color=SLATE_DARK)


def render_field_block(parent_cell: Any, fields: list[SurveyField]) -> None:
    """Render a sequence of fields, switching between grid and full-width rows."""
    queue: list[SurveyField] = []

    def flush() -> None:
        if queue:
            render_short_field_grid(parent_cell, list(queue))
            queue.clear()

    for field_item in fields:
        if is_full_width_field(field_item) or not is_short_field(field_item):
            flush()
            render_full_width_field(parent_cell, field_item)
        else:
            queue.append(field_item)
    flush()


# ---------------------------------------------------------------------------
# Repeat rendering
# ---------------------------------------------------------------------------


def flatten_repeat_fields(repeat: SurveyRepeat) -> list[tuple[str, SurveyField | None]]:
    rows: list[tuple[str, SurveyField | None]] = []

    def walk_group(group: SurveyGroup, prefix: str) -> None:
        group_label = group.label or group.name
        section_title = f"{prefix} - {group_label}" if prefix and group_label else (prefix or group_label)
        if group.fields:
            rows.append((section_title or group_label, None))
            for field_item in group.fields:
                rows.append((field_item.label, field_item))
        for child in group.groups:
            walk_group(child, section_title)
        for nested in group.repeats:
            rows.append((f"{section_title} - {nested.label or nested.name}", None))

    if repeat.fields:
        rows.append(("Details", None))
        for field_item in repeat.fields:
            rows.append((field_item.label, field_item))
    for child_group in repeat.groups:
        walk_group(child_group, "")

    return rows


def render_repeat_table(doc: Document, repeat: SurveyRepeat) -> None:
    rows = flatten_repeat_fields(repeat)
    label = repeat.label or repeat.name

    intro = doc.add_paragraph()
    clear_paragraph_spacing(intro)
    intro.paragraph_format.space_before = Pt(4)
    add_text(intro, f"REPEAT RECORDS - {label.upper()}", bold=True, size=9, color=NAVY)

    if not rows:
        empty = doc.add_paragraph()
        clear_paragraph_spacing(empty)
        add_text(empty, f"${{# {repeat.name}}}${{{repeat.name}}}${{/}}", size=8, color=SLATE_MID)
        return

    table = doc.add_table(rows=len(rows) + 4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=DIVIDER_HEX, size="4")

    field_col_width = 2.6
    response_col_width = PAGE_WIDTH_IN - field_col_width

    open_cells = table.rows[0].cells
    open_merged = open_cells[0].merge(open_cells[1])
    set_cell_width(open_merged, PAGE_WIDTH_IN)
    set_cell_shading(open_merged, AMBER_HEX)
    set_cell_margins(open_merged, top=70, bottom=70, start=140, end=140)
    open_paragraph = open_merged.paragraphs[0]
    clear_paragraph_spacing(open_paragraph)
    add_text(open_paragraph, f"${{# {repeat.name}}}", bold=True, size=8, color=NAVY, font=MONO_FONT)

    record_cells = table.rows[1].cells
    record_merged = record_cells[0].merge(record_cells[1])
    set_cell_width(record_merged, PAGE_WIDTH_IN)
    set_cell_shading(record_merged, NAVY_HEX)
    set_cell_margins(record_merged, top=80, bottom=80, start=140, end=140)
    record_paragraph = record_merged.paragraphs[0]
    clear_paragraph_spacing(record_paragraph)
    add_text(record_paragraph, "RECORD ", bold=True, size=10, color=WHITE)
    add_text(record_paragraph, "${$position}", bold=True, size=10, color=AMBER, font=MONO_FONT)
    add_text(record_paragraph, f"   |   {label}", bold=False, size=9, color=WHITE)

    header_field_cell, header_response_cell = table.rows[2].cells
    set_cell_width(header_field_cell, field_col_width)
    set_cell_width(header_response_cell, response_col_width)
    set_cell_shading(header_field_cell, SUBSECTION_HEX)
    set_cell_shading(header_response_cell, SUBSECTION_HEX)
    set_cell_margins(header_field_cell, top=60, bottom=60, start=140, end=140)
    set_cell_margins(header_response_cell, top=60, bottom=60, start=140, end=140)
    header_field_cell.paragraphs[0].text = ""
    header_response_cell.paragraphs[0].text = ""
    add_text(header_field_cell.paragraphs[0], "Field", bold=True, size=8, color=NAVY)
    add_text(header_response_cell.paragraphs[0], "Response", bold=True, size=8, color=NAVY)

    for index, (label_text, field_item) in enumerate(rows, start=3):
        row_cells = table.rows[index].cells
        label_cell, value_cell = row_cells
        set_cell_width(label_cell, field_col_width)
        set_cell_width(value_cell, response_col_width)
        set_cell_margins(label_cell, top=60, bottom=60, start=140, end=140)
        set_cell_margins(value_cell, top=60, bottom=60, start=140, end=140)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_cell.paragraphs[0].text = ""
        value_cell.paragraphs[0].text = ""

        if field_item is None:
            merged = label_cell.merge(value_cell)
            set_cell_shading(merged, CREAM_HEX)
            merged_paragraph = merged.paragraphs[0]
            clear_paragraph_spacing(merged_paragraph)
            add_text(merged_paragraph, label_text, bold=True, size=8, color=NAVY)
            continue

        label_paragraph = label_cell.paragraphs[0]
        clear_paragraph_spacing(label_paragraph)
        add_text(label_paragraph, label_text, bold=True, size=8, color=SLATE_DARK)

        value_paragraph = value_cell.paragraphs[0]
        clear_paragraph_spacing(value_paragraph)
        if is_short_field(field_item):
            _, response_text = short_field_block_text(field_item)
        else:
            response_text = field_response(field_item)
        add_text(value_paragraph, response_text, size=9, color=SLATE_DARK)

    close_cells = table.rows[-1].cells
    close_merged = close_cells[0].merge(close_cells[1])
    set_cell_width(close_merged, PAGE_WIDTH_IN)
    set_cell_shading(close_merged, AMBER_HEX)
    set_cell_margins(close_merged, top=70, bottom=70, start=140, end=140)
    close_paragraph = close_merged.paragraphs[0]
    clear_paragraph_spacing(close_paragraph)
    add_text(close_paragraph, "${/}", bold=True, size=8, color=NAVY, font=MONO_FONT)

    spacer = doc.add_paragraph()
    clear_paragraph_spacing(spacer)
    spacer.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Cover & rendering pipeline
# ---------------------------------------------------------------------------


def add_cover_page(doc: Document, form: SurveyForm) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(eyebrow, "ARCGIS ONLINE FEATURE REPORT", bold=True, size=9, color=AMBER)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(form.title.upper())
    title_run.bold = True
    title_run.font.name = BODY_FONT
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = NAVY

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(rule, AMBER_HEX)
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(8)
    add_text(rule, " ", size=2, color=AMBER)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(
        summary,
        "Inspection dossier template - auto-generated from the Survey123 Connect XLSForm. "
        "Yes/No questions render as live checkbox glyphs and repeat groups print as iterating tables.",
        italic=True,
        size=9,
        color=SLATE_MID,
    )
    doc.add_paragraph()

    meta_rows: list[tuple[str, str]] = [
        ("Site", "${loc_city}, ${loc_state}  (${loc_id})"),
        ("Facility Type", "${fac_type}"),
        ("JCN / JON", "${jcn} / ${jon}"),
        ("Address", "${address}"),
        ("County / Service Area", "${county} / ${service_area}"),
        ("Report Generated", '${$feature | getValue:"" | format:"MM/DD/YYYY hh:mm a"}'),
        ("Survey Version", "${survey_version}"),
    ]
    if form.version:
        meta_rows.append(("XLSForm Version", form.version))
    if form.form_id:
        meta_rows.append(("Form ID", form.form_id))

    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=DIVIDER_HEX, size="4")

    for index, (label, value) in enumerate(meta_rows):
        label_cell, value_cell = table.rows[index].cells
        set_cell_width(label_cell, 2.2)
        set_cell_width(value_cell, PAGE_WIDTH_IN - 2.2)
        set_cell_shading(label_cell, SUBSECTION_HEX)
        set_cell_margins(label_cell, top=80, bottom=80, start=160, end=160)
        set_cell_margins(value_cell, top=80, bottom=80, start=160, end=160)
        label_cell.paragraphs[0].text = ""
        value_cell.paragraphs[0].text = ""
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(label_cell.paragraphs[0], label, bold=True, size=8, color=NAVY)
        add_text(value_cell.paragraphs[0], value, size=9, color=SLATE_DARK)

    spacer = doc.add_paragraph()
    clear_paragraph_spacing(spacer)
    spacer.paragraph_format.space_before = Pt(8)
    add_text(
        spacer,
        "Filled checkboxes (\u2612) represent the recorded answer for each question. "
        "Empty boxes (\u2610) indicate an unselected option.",
        italic=True,
        size=8,
        color=SLATE_MID,
    )


def render_subsection(doc: Document, number: str, group: SurveyGroup) -> None:
    title = group.label or group.name
    if title:
        add_subsection_banner(doc, number, title)
    content_cell = open_sidebar_frame(doc)
    render_field_block(content_cell, group.fields)
    for repeat in group.repeats:
        render_repeat_table(doc, repeat)
    for child_index, child_group in enumerate(group.groups, start=1):
        render_subsection(doc, f"{number}.{child_index}", child_group)
    spacer = doc.add_paragraph()
    clear_paragraph_spacing(spacer)
    spacer.paragraph_format.space_after = Pt(2)


def render_section(doc: Document, number: int, group: SurveyGroup) -> None:
    title = group.label or group.name
    add_section_banner(doc, f"{number}.0", title or f"Section {number}")
    if group.fields:
        content_cell = open_sidebar_frame(doc)
        render_field_block(content_cell, group.fields)
    for sub_index, child_group in enumerate(group.groups, start=1):
        render_subsection(doc, f"{number}.{sub_index}", child_group)
    for repeat in group.repeats:
        render_repeat_table(doc, repeat)
    spacer = doc.add_paragraph()
    clear_paragraph_spacing(spacer)


def render_form(doc: Document, form: SurveyForm) -> None:
    section_index = 0
    orphan_fields: list[SurveyField] = []
    orphan_repeats: list[SurveyRepeat] = []
    for item in form.items:
        if isinstance(item, SurveyGroup):
            section_index += 1
            render_section(doc, section_index, item)
        elif isinstance(item, SurveyRepeat):
            orphan_repeats.append(item)
        else:
            orphan_fields.append(item)

    if orphan_fields or orphan_repeats:
        section_index += 1
        add_section_banner(doc, f"{section_index}.0", "Additional Fields")
        if orphan_fields:
            content_cell = open_sidebar_frame(doc)
            render_field_block(content_cell, orphan_fields)
        for repeat in orphan_repeats:
            render_repeat_table(doc, repeat)


def add_signature_block(doc: Document) -> None:
    spacer = doc.add_paragraph()
    clear_paragraph_spacing(spacer)
    spacer.paragraph_format.space_before = Pt(8)

    add_section_banner(doc, "S.0", "Sign-off")
    content_cell = open_sidebar_frame(doc)
    table = content_cell.add_table(rows=2, cols=2)
    set_table_borders(table, color=DIVIDER_HEX, size="4")
    width = (PAGE_WIDTH_IN - SIDEBAR_WIDTH_IN - 0.25) / 2

    captions = [
        ("Inspector", "${creator}"),
        ("Date Completed", '${$feature | getValue:"" | format:"MM/DD/YYYY"}'),
        ("Reviewer", " "),
        ("Notes", " "),
    ]

    for index, (label, value) in enumerate(captions):
        row_index, col_index = divmod(index, 2)
        cell = table.rows[row_index].cells[col_index]
        cell.text = ""
        set_cell_width(cell, width)
        set_cell_margins(cell, top=100, bottom=160, start=140, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        label_paragraph = cell.paragraphs[0]
        clear_paragraph_spacing(label_paragraph)
        add_text(label_paragraph, label, bold=True, size=8, color=NAVY)
        response_paragraph = cell.add_paragraph()
        clear_paragraph_spacing(response_paragraph)
        add_text(response_paragraph, value, size=10, color=SLATE_DARK)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def build_document(form: SurveyForm) -> Document:
    document = Document()
    add_cover_page(document, form)
    document.add_page_break()
    render_form(document, form)
    add_signature_block(document)
    return document


def unique_output_path(path: Path) -> Path:
    """Return a non-conflicting path: appends `-v2`, `-v3`, ... if needed."""
    if not path.exists():
        return path
    stem = path.stem
    parent = path.parent
    suffix = path.suffix
    counter = 2
    while True:
        candidate = parent / f"{stem}-v{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def generate_report_template(xlsx_path: Path, output_path: Path) -> Path:
    form = parse_xlsform(xlsx_path)
    document = build_document(form)
    resolved_output = unique_output_path(output_path)
    resolved_output.parent.mkdir(parents=True, exist_ok=True)
    document.save(resolved_output)
    return resolved_output


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Generate an LPGBS-style inspection dossier feature report template "
            "(.docx) from a Survey123 Connect XLSForm. Existing files are never "
            "overwritten - a versioned filename is used if the target already exists."
        )
    )
    parser.add_argument(
        "--xlsx",
        required=True,
        type=Path,
        help="Path to the Survey123 Connect XLSForm (.xlsx).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("outputs/lpgbs-inspection-report.docx"),
        help=(
            "Output .docx path. Defaults to outputs/lpgbs-inspection-report.docx. "
            "If the file already exists the generator appends -v2, -v3, ... to "
            "avoid overwriting work."
        ),
    )
    args = parser.parse_args()

    if not args.xlsx.exists():
        raise SystemExit(f"XLSForm not found: {args.xlsx}")

    output = generate_report_template(args.xlsx, args.output)
    print(f"Wrote report template: {output}")


if __name__ == "__main__":
    main()
